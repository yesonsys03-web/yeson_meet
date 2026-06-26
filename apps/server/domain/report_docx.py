# === ANCHOR: REPORT_DOCX_START ===
"""Word (.docx) report generation for completed meeting sessions (S3)."""
from __future__ import annotations

import io
import re
from itertools import groupby

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from apps.server.db.models import Session, Utterance
from apps.server.domain.reports import _hms, _speaker_label, _to_local, merge_continuation_utterances

# Korean fonts to attempt; Word will fall back gracefully if not installed.
_KO_FONT = "맑은 고딕"
_EN_FONT = "Calibri"


def _set_run_font(run, font_name: str) -> None:
    """Apply both the ASCII and East-Asian font slot for a run."""
    run.font.name = font_name
    # Force the East Asian (hAnsi / eastAsia) slot so Korean characters render.
    r = run._r  # noqa: SLF001  (docx internal)
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _add_inline_md(paragraph, text: str) -> None:
    """Add inline markdown (**bold**, `code`) as styled runs to a paragraph.

    Splits on ``**...**`` markers; strips backtick code spans to plain text.
    """
    # Split on **...** — odd indices are bold, even indices are plain.
    parts = re.split(r"\*\*(.+?)\*\*", text)
    bold = False
    for part in parts:
        # Strip `code` backticks to plain text (keep content).
        clean = re.sub(r"`([^`]+)`", r"\1", part)
        if clean:
            run = paragraph.add_run(clean)
            run.bold = bold
            _set_run_font(run, _KO_FONT)
        bold = not bold


def _summary_md_to_docx(doc: Document, summary: str) -> None:
    """Render the small Markdown subset used in LLM summaries as Word content.

    Handles: ## headings (level = min(hash_count+1, 4)), - /* bullet lists,
    --- separators (skipped), **bold** and `code` inline.  Blank lines skipped.
    """
    for raw in summary.splitlines():
        line = raw.strip()
        if not line:
            continue
        # Separator lines — skip (no Word equivalent needed).
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line):
            continue
        # Heading.
        heading_m = re.match(r"(#{1,6})\s+(.*)", line)
        if heading_m:
            level = min(len(heading_m.group(1)) + 1, 4)
            h = doc.add_heading(heading_m.group(2), level=level)
            for run in h.runs:
                _set_run_font(run, _KO_FONT)
            continue
        # Bullet list item.
        bullet_m = re.match(r"[-*]\s+(.*)", line)
        if bullet_m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_md(p, bullet_m.group(1))
            continue
        # Normal paragraph.
        p = doc.add_paragraph()
        _add_inline_md(p, line)


# === ANCHOR: REPORT_DOCX_BUILD_START ===
def build_session_report_docx(
    meeting: Session,
    utterances: list[Utterance],
    summary: str | None = None,
) -> bytes:
    """Build a Word (.docx) report for one meeting and return its bytes.

    Applies the same speaker-grouping and KO-first rules as S1/S2:
    - Consecutive utterances by the same speaker are merged under one header.
    - Block header: HH:MM:SS + speaker label (None → 발화자 미상).
    - KO body text first, EN body text below in smaller/greyed style.

    *summary* is an optional LLM-generated Korean summary (S6).  When provided
    it is inserted as a ``## 요약`` heading section before ``Utterances``.
    """
    doc = Document()

    # --- Title (Heading 1) ---
    title_para = doc.add_heading(meeting.title or "(제목 없음)", level=1)
    for run in title_para.runs:
        _set_run_font(run, _KO_FONT)

    # --- Summary statistics ---
    started = _to_local(meeting.started_at).isoformat() if meeting.started_at else "N/A"
    ended = _to_local(meeting.ended_at).isoformat() if meeting.ended_at else "N/A"
    meta_lines = [
        f"Session: {meeting.external_id}",
        f"Status: {meeting.status}",
        f"Started: {started}",
        f"Ended: {ended}",
    ]
    if meeting.client_label:
        meta_lines.append(f"Client: {meeting.client_label}")

    if utterances:
        speakers = sorted({_speaker_label(r.speaker) for r in utterances if r.speaker})
        duration = (
            f"{_hms(utterances[0].started_at)} – {_hms(utterances[-1].ended_at)}"
        )
        meta_lines.extend(
            [
                f"참여 화자: {', '.join(speakers) if speakers else '없음'}",
                f"시간 범위: {duration}",
            ]
        )

    for line in meta_lines:
        p = doc.add_paragraph(line, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_run_font(run, _KO_FONT)

    # --- LLM summary section (S6) ---
    if summary:
        h2_sum = doc.add_heading("요약", level=2)
        for run in h2_sum.runs:
            _set_run_font(run, _KO_FONT)
        _summary_md_to_docx(doc, summary)

    # --- Utterances heading ---
    h2 = doc.add_heading("Utterances", level=2)
    for run in h2.runs:
        _set_run_font(run, _KO_FONT)

    # --- Utterance blocks ---
    if not utterances:
        p = doc.add_paragraph("(No utterances recorded.)")
        for run in p.runs:
            run.font.italic = True
            _set_run_font(run, _KO_FONT)
    else:
        # Group consecutive utterances by speaker (None treated as distinct key).
        # Merge continuation rows (empty text_en, same speaker) first.
        for speaker_key, group in groupby(merge_continuation_utterances(utterances), key=lambda r: r.speaker):
            group_rows = list(group)
            label = _speaker_label(speaker_key)
            first = group_rows[0]

            # Block header: HH:MM:SS + speaker label
            header_text = f"{_hms(first.started_at)}  {label}"
            h3 = doc.add_heading(header_text, level=3)
            for run in h3.runs:
                _set_run_font(run, _KO_FONT)

            for row in group_rows:
                # KO — primary body, full-weight
                ko_text = row.text_ko or ""
                p_ko = doc.add_paragraph()
                run_ko = p_ko.add_run(ko_text)
                run_ko.font.size = Pt(12)
                run_ko.bold = True
                _set_run_font(run_ko, _KO_FONT)
                p_ko.paragraph_format.space_after = Pt(1)

                # EN — secondary body, smaller + greyed
                en_text = row.text_en or ""
                p_en = doc.add_paragraph()
                run_en = p_en.add_run(en_text)
                run_en.font.size = Pt(10)
                run_en.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                _set_run_font(run_en, _EN_FONT)
                p_en.paragraph_format.space_after = Pt(6)

    # --- Serialise to bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
# === ANCHOR: REPORT_DOCX_BUILD_END ===


# === ANCHOR: REPORT_DOCX_SUMMARY_BUILD_START ===
def build_summary_docx(meeting: Session, summary: str) -> bytes:
    """Build a minimal Word (.docx) document for a standalone summary.

    Reuses the Korean-font handling pattern. Contains a title Heading and the
    summary body (each non-blank line as a paragraph).
    """
    doc = Document()

    title_para = doc.add_heading(f"요약 — {meeting.title or '(제목 없음)'}", level=1)
    for run in title_para.runs:
        _set_run_font(run, _KO_FONT)

    _summary_md_to_docx(doc, summary)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
# === ANCHOR: REPORT_DOCX_SUMMARY_BUILD_END ===
# === ANCHOR: REPORT_DOCX_END ===
