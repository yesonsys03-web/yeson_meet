# === ANCHOR: TEST_REPORT_SUMMARY_START ===
"""Tests for S6: report_summary.generate_summary() and find_summary_cli().

Run without conftest (no DB required):
    uv run pytest apps/server/tests/test_report_summary.py -v --noconftest
"""
from __future__ import annotations

import subprocess
from datetime import datetime, timezone
from types import SimpleNamespace
from unittest.mock import patch

import pytest

from apps.server.domain.report_summary import find_summary_cli, generate_summary


# ---------------------------------------------------------------------------
# Helpers — lightweight stubs
# ---------------------------------------------------------------------------

def _meeting(title: str = "Test Meeting") -> SimpleNamespace:
    base = datetime(2026, 6, 24, 9, 0, 0, tzinfo=timezone.utc)
    return SimpleNamespace(
        title=title,
        external_id="test-session-001",
        status="ended",
        started_at=base,
        ended_at=datetime(2026, 6, 24, 9, 30, 0, tzinfo=timezone.utc),
        client_label=None,
    )


def _utt(
    speaker: str | None = "Alice",
    text_ko: str = "안녕하세요.",
    text_en: str = "Hello.",
    seq: int = 1,
) -> SimpleNamespace:
    base = datetime(2026, 6, 24, 9, 5, 0, tzinfo=timezone.utc)
    return SimpleNamespace(
        seq=seq,
        speaker=speaker,
        text_ko=text_ko,
        text_en=text_en,
        started_at=base,
        ended_at=base,
    )


# ---------------------------------------------------------------------------
# (i) CLI found + subprocess returns stdout → generate_summary returns text
# ---------------------------------------------------------------------------

def test_generate_summary_returns_stdout_when_cli_succeeds(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("YESON_REPORT_SUMMARY", raising=False)

    fake_result = SimpleNamespace(returncode=0, stdout="회의 요약 내용입니다.\n", stderr="")

    with (
        patch(
            "apps.server.domain.report_summary.find_summary_cli",
            return_value=("claude", ["claude", "-p"]),
        ),
        patch("subprocess.run", return_value=fake_result) as mock_run,
    ):
        result = generate_summary(_meeting(), [_utt()])

    assert result == "회의 요약 내용입니다."
    mock_run.assert_called_once()
    call_args = mock_run.call_args[0][0]
    assert call_args[:2] == ["claude", "-p"]


# ---------------------------------------------------------------------------
# (ii) find_summary_cli returns None → generate_summary returns None, no exception
# ---------------------------------------------------------------------------

def test_generate_summary_returns_none_when_no_cli(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("YESON_REPORT_SUMMARY", raising=False)

    with patch(
        "apps.server.domain.report_summary.find_summary_cli",
        return_value=None,
    ):
        result = generate_summary(_meeting(), [_utt()])

    assert result is None  # no exception raised


# ---------------------------------------------------------------------------
# (iii-a) returncode != 0 → None, no exception
# ---------------------------------------------------------------------------

def test_generate_summary_returns_none_on_nonzero_returncode(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("YESON_REPORT_SUMMARY", raising=False)

    fake_result = SimpleNamespace(returncode=1, stdout="", stderr="error occurred")

    with (
        patch(
            "apps.server.domain.report_summary.find_summary_cli",
            return_value=("claude", ["claude", "-p"]),
        ),
        patch("subprocess.run", return_value=fake_result),
    ):
        result = generate_summary(_meeting(), [_utt()])

    assert result is None


# ---------------------------------------------------------------------------
# (iii-b) TimeoutExpired → None, no exception
# ---------------------------------------------------------------------------

def test_generate_summary_returns_none_on_timeout(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("YESON_REPORT_SUMMARY", raising=False)

    with (
        patch(
            "apps.server.domain.report_summary.find_summary_cli",
            return_value=("claude", ["claude", "-p"]),
        ),
        patch(
            "subprocess.run",
            side_effect=subprocess.TimeoutExpired(cmd=["claude", "-p", "x"], timeout=120),
        ),
    ):
        result = generate_summary(_meeting(), [_utt()])

    assert result is None


# ---------------------------------------------------------------------------
# (iv) YESON_REPORT_SUMMARY=0 → None (feature disabled)
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("env_val", ["0", "false", "False", "off", "OFF"])
def test_generate_summary_disabled_by_env(
    monkeypatch: pytest.MonkeyPatch,
    env_val: str,
) -> None:
    monkeypatch.setenv("YESON_REPORT_SUMMARY", env_val)

    with patch("subprocess.run") as mock_run:
        result = generate_summary(_meeting(), [_utt()])

    assert result is None
    mock_run.assert_not_called()  # subprocess must never be invoked when disabled


# ---------------------------------------------------------------------------
# (v) Empty utterances → None (nothing to summarise)
# ---------------------------------------------------------------------------

def test_generate_summary_returns_none_for_empty_utterances(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("YESON_REPORT_SUMMARY", raising=False)

    with (
        patch(
            "apps.server.domain.report_summary.find_summary_cli",
            return_value=("claude", ["claude", "-p"]),
        ),
        patch("subprocess.run") as mock_run,
    ):
        result = generate_summary(_meeting(), [])

    assert result is None
    mock_run.assert_not_called()


# ---------------------------------------------------------------------------
# (vi-a) build_session_report with summary → ## 요약 before ## Utterances
# ---------------------------------------------------------------------------

def test_build_session_report_summary_injected_before_utterances() -> None:
    from apps.server.domain.reports import build_session_report

    meeting = _meeting()
    utts = [_utt()]
    result = build_session_report(meeting, utts, summary="요약내용입니다.")

    assert "## 요약" in result
    assert "요약내용입니다." in result
    assert "## Utterances" in result

    summary_idx = result.index("## 요약")
    utterances_idx = result.index("## Utterances")
    assert summary_idx < utterances_idx, "## 요약 must appear before ## Utterances"


# (vi-b) build_session_report without summary → no 요약 section
def test_build_session_report_no_summary_section_when_none() -> None:
    from apps.server.domain.reports import build_session_report

    result = build_session_report(_meeting(), [_utt()])
    assert "## 요약" not in result


# ---------------------------------------------------------------------------
# (vi-c) build_session_report_html with summary → 요약 block before utterances
# ---------------------------------------------------------------------------

def test_build_session_report_html_summary_injected() -> None:
    from apps.server.domain.report_html import build_session_report_html

    meeting = _meeting()
    utts = [_utt()]
    result = build_session_report_html(meeting, utts, summary="HTML 요약 내용")

    assert "요약" in result
    assert "HTML 요약 내용" in result
    # summary section must appear before speaker blocks
    summary_idx = result.index("요약")
    speaker_block_idx = result.index('class="speaker-block"')
    assert summary_idx < speaker_block_idx


# ---------------------------------------------------------------------------
# (vi-d) build_session_report_docx with summary → 요약 heading in document
# ---------------------------------------------------------------------------

def test_build_session_report_docx_summary_injected() -> None:
    import io
    from docx import Document
    from apps.server.domain.report_docx import build_session_report_docx

    meeting = _meeting()
    utts = [_utt()]
    raw = build_session_report_docx(meeting, utts, summary="DOCX 요약 내용")

    doc = Document(io.BytesIO(raw))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "요약" in all_text
    assert "DOCX 요약 내용" in all_text

    # 요약 heading must appear before Utterances heading in paragraph order
    texts = [p.text for p in doc.paragraphs]
    yoyak_idx = next((i for i, t in enumerate(texts) if "요약" in t), None)
    utt_idx = next((i for i, t in enumerate(texts) if t == "Utterances"), None)
    assert yoyak_idx is not None
    assert utt_idx is not None
    assert yoyak_idx < utt_idx, "요약 must appear before Utterances in docx"


# ---------------------------------------------------------------------------
# find_summary_cli: prefers claude over codex
# ---------------------------------------------------------------------------

def test_find_summary_cli_prefers_claude_over_codex() -> None:
    with patch(
        "apps.server.domain.report_summary.shutil.which",
        side_effect=lambda name: "/usr/local/bin/claude" if name == "claude" else None,
    ):
        result = find_summary_cli()

    assert result is not None
    name, args = result
    assert name == "claude"
    # argv is the resolved absolute path (so Windows .cmd shims run correctly);
    # the prompt is delivered on stdin, not appended here.
    assert args == ["/usr/local/bin/claude", "-p"]


def test_find_summary_cli_falls_back_to_codex() -> None:
    with patch(
        "apps.server.domain.report_summary.shutil.which",
        side_effect=lambda name: "/usr/local/bin/codex" if name == "codex" else None,
    ):
        result = find_summary_cli()

    assert result is not None
    name, args = result
    assert name == "codex"
    assert args == ["/usr/local/bin/codex", "exec"]


def test_find_summary_cli_returns_none_when_neither_found() -> None:
    with patch(
        "apps.server.domain.report_summary.shutil.which",
        return_value=None,
    ):
        result = find_summary_cli()

    assert result is None


# ---------------------------------------------------------------------------
# Standalone summary builders (multi-format summary export)
# ---------------------------------------------------------------------------

def test_build_summary_html_contains_title_and_body() -> None:
    from apps.server.domain.report_html import build_summary_html

    result = build_summary_html(_meeting(), "요약 본문 한 줄")

    assert "<!DOCTYPE html>" in result
    assert "요약 — Test Meeting" in result
    assert "요약 본문 한 줄" in result


def test_build_summary_html_escapes_body() -> None:
    from apps.server.domain.report_html import build_summary_html

    result = build_summary_html(_meeting(), "<script>x</script>")
    assert "<script>x</script>" not in result
    assert "&lt;script&gt;" in result


def test_build_summary_docx_contains_title_and_body() -> None:
    import io
    from docx import Document
    from apps.server.domain.report_docx import build_summary_docx

    raw = build_summary_docx(_meeting(), "DOCX 요약 본문")
    doc = Document(io.BytesIO(raw))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "요약 — Test Meeting" in all_text
    assert "DOCX 요약 본문" in all_text
# === ANCHOR: TEST_REPORT_SUMMARY_END ===
